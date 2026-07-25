# -*- coding: utf-8 -*-
import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

# 1. Asegurar directorios
os.makedirs('estate_management/data/templates', exist_ok=True)
os.makedirs('estate_management/static/src/img', exist_ok=True)

wm_path = 'estate_management/static/src/img/marcaagua.jpeg'
if not os.path.exists(wm_path):
    wm_path = 'estate_management/data/marcaagua.jpeg'
if not os.path.exists(wm_path):
    wm_path = '/home/justin/Documentos/Tesis/marcaagua.jpeg'

# 2. Función para configurar márgenes de izquierda/derecha EXACTOS a los docx originales, y márgenes seguros superior/inferior + marca de agua A4
def apply_watermark_and_margins(doc, img_path):
    for section in doc.sections:
        # Asegurar formato A4 exacto
        if not section.page_width or section.page_width.inches < 7.0:
            section.page_width = Inches(8.268)   # A4 width
            section.page_height = Inches(11.693) # A4 height
        
        # Mantenemos intactos los márgenes originales superior, inferior, izquierdo y derecho del documento proporcionado por el usuario
        if not os.path.exists(img_path):
            continue
        
        header = section.header
        if not header.paragraphs:
            header.add_paragraph()
        p = header.paragraphs[0]
        run = p.add_run()
        
        # Insertar imagen con las dimensiones exactas de la página (A4 completa)
        pic = run.add_picture(img_path, width=section.page_width, height=section.page_height)
        inline = pic._inline
        drawing = inline.getparent()
        
        cx = inline.find(etree.QName("http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing", "extent")).get("cx")
        cy = inline.find(etree.QName("http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing", "extent")).get("cy")
        
        # Convertir a anchor flotante posicionado en (0,0) respecto a la página y detrás del texto (behindDoc="1")
        anchor_xml = f"""
        <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                   xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                   xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
                   xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                   distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">
          <wp:simplePos x="0" y="0"/>
          <wp:positionH relativeFrom="page">
            <wp:align>center</wp:align>
          </wp:positionH>
          <wp:positionV relativeFrom="page">
            <wp:align>center</wp:align>
          </wp:positionV>
          <wp:extent cx="{cx}" cy="{cy}"/>
          <wp:effectExtent l="0" t="0" r="0" b="0"/>
          <wp:wrapNone/>
          <wp:docPr id="999" name="MarcaAguaOficial"/>
          <wp:cNvGraphicFramePr>
            <a:graphicFrameLocks noChangeAspect="1"/>
          </wp:cNvGraphicFramePr>
        </wp:anchor>
        """
        anchor = etree.fromstring(anchor_xml)
        graphic = inline.find(etree.QName("http://schemas.openxmlformats.org/drawingml/2006/main", "graphic"))
        if graphic is not None:
            anchor.append(graphic)
            drawing.replace(inline, anchor)

mapping = {
    'CONTRATO DE CORRETAJE CON EXCLUSIVIDAD ACTUAL.docx': 'tpl_exclusive_owner.docx',
    'CONTRATO DE CORRETAJE CON EXCLUSIVIDAD APODERADO.docx': 'tpl_exclusive_proxy.docx',
    'CONTRATO DE CORRETAJE SIN EXCLUSIVIDAD.docx': 'tpl_non_exclusive_owner.docx',
    'CONTRATO DE CORRETAJE SIN EXCLUSIVIDAD apoderado.docx': 'tpl_non_exclusive_proxy.docx',
}

# 3. Procesar párrafos conservando EXACTAMENTE el tipo de letra original (Helvetica) y el tamaño original de cada párrafo (ej: 8 pt)
def process_paragraph_keep_exact_font_and_size(p, is_proxy=False):
    text = p.text
    if not text.strip():
        return
    
    new_text = text

    # 1. Header date
    if 'En la ciudad de' in new_text and 'días del mes de' in new_text:
        new_text = re.sub(r'En la ciudad de[^,]+,\s*a los[^d]+días del mes de[^d]+del año[^,]+,',
                          'En la ciudad de {{ ciudad }}, a los {{ dia }} días del mes de {{ mes }} del año {{ anio }},', new_text)
        if '{{ ciudad }}' not in new_text:
            new_text = re.sub(r'En la ciudad de.*?del año.*?,',
                              'En la ciudad de {{ ciudad }}, a los {{ dia }} días del mes de {{ mes }} del año {{ anio }},', new_text)

    # 2. Comparecientes
    if 'PRIMERA: COMPARECIENTES' in new_text:
        if is_proxy:
            new_text = re.sub(r'CLIENTE U OFERENTE.*?con clave catastral número.*?[,;]\s*y (?:por|de la) otra parte',
                              'CLIENTE U OFERENTE, {{ cliente }}, con número de cédula {{ cedula_cliente }}, de estado civil {{ estado_civil }}, en calidad de apoderado/a de {{ propietario }}, cuyo poder se adjunta como documento habilitante; propietario del inmueble ubicado en la ciudad de {{ ciudad }}, sector {{ sector }}, calle {{ calle }}, número {{ numero }}, con clave catastral número {{ clave_catastral }}, y por otra parte', new_text, flags=re.DOTALL)
        else:
            new_text = re.sub(r'CLIENTE U OFERENTE.*?con clave catastral número.*?[,;]\s*y (?:por|de la) otra parte',
                              'CLIENTE U OFERENTE, {{ cliente }}, con número de cédula {{ cedula_cliente }}, de estado civil {{ estado_civil }}, propietario del inmueble ubicado en la ciudad de {{ ciudad }}, sector {{ sector }}, calle {{ calle }}, número {{ numero }}, con clave catastral número {{ clave_catastral }}, y por otra parte', new_text, flags=re.DOTALL)

    # 3. Plazo
    if 'PLAZO' in new_text and 'tendrá una duración' in new_text:
        new_text = re.sub(r'tendrá una duración.*?días contados',
                          'tendrá una duración de {{ plazo_dias }} días contados', new_text)

    # 4. Precio de Venta
    if 'PRECIO DE VENTA' in new_text and 'fija en la cantidad' in new_text:
        new_text = re.sub(r'fija en la cantidad.*?Dólares',
                          'fija en la cantidad de ${{ precio }} ({{ precio_letras }}) Dólares', new_text)

    # 5. Honorarios
    if 'HONORARIOS' in new_text and 'profesionales al CORREDOR' in new_text:
        new_text = re.sub(r'profesionales al CORREDOR,\s*el.*?por ciento\)',
                          'profesionales al CORREDOR, el {{ honorarios }} % ({{ honorarios_letras }} por ciento)', new_text)

    # 6. Signature blocks CI (conservando tabs y espacios originales para que CI y cédula del cliente queden perfectamente alineados)
    if 'CI:' in new_text and 'cedula_' not in new_text:
        parts = new_text.split('CI:')
        cedula_marker = '{{ cedula_apoderado_firma }}' if is_proxy else '{{ cedula_cliente_firma }}'
        if len(parts) >= 3:
            # Mantener exactos los espacios/tabs originales entre el primer CI: y el segundo CI:
            new_text = parts[0] + 'CI:' + parts[1] + 'CI: ' + cedula_marker + parts[2]
        elif len(parts) == 2:
            new_text = parts[0] + 'CI: ' + cedula_marker + parts[1]

    # Detectar el tamaño y fuente original del párrafo en el documento docx proporcionado
    detected_font = 'Helvetica'
    detected_size = Pt(8.0) # tamaño por defecto en el docx original
    for r in p.runs:
        if r.text.strip():
            if r.font.name:
                detected_font = r.font.name
            if r.font.size:
                detected_size = r.font.size
            break
            
    is_header_line = new_text.strip().isupper() and len(new_text.strip()) > 3 and '{{' not in new_text
    is_contract_title = new_text.strip().startswith('CONTRATO DE CORRETAJE')
    if is_contract_title:
        detected_size = Pt(11.0)

    # Reconstruir runs preservando exactamente la fuente original ('Helvetica') y el tamaño original ('detected_size', ej: 8 pt)
    if new_text != text or '{{' in new_text or any(new_text.startswith(k) for k in ('PRIMERA:', 'SEGUNDA:', 'TERCERA:', 'CUARTA:', 'QUINTA:', 'SEXTA:', 'SÉPTIMA:', 'OCTAVA:', 'NOVENA:', 'DÉCIMA:')):
        p.clear()
        pattern = r'(\{\{\s*[\w_]+\s*\}\}|PRIMERA:|SEGUNDA:|TERCERA:|CUARTA:|QUINTA:|SEXTA:|SÉPTIMA:|OCTAVA:|NOVENA:|DÉCIMA:|INMOBI Bienes Raíces|EL CLIENTE U OFERENTE|CLIENTE U OFERENTE|CORREDOR\.)'
        tokens = re.split(pattern, new_text)
        for token in tokens:
            if not token:
                continue
            r = p.add_run(token)
            r.font.name = detected_font
            r.font.size = detected_size
            if is_header_line or re.match(pattern, token):
                r.bold = True
            else:
                r.bold = False
    else:
        # Asegurar que los párrafos no modificados también conserven su fuente y tamaño exacto original
        for r in p.runs:
            if not r.font.name:
                r.font.name = detected_font
            if not r.font.size:
                r.font.size = detected_size

for src, dst in mapping.items():
    if not os.path.exists(src):
        print(f"Saltando {src}, no se encontró.")
        continue
    doc = docx.Document(src)
    is_proxy = 'apoderado' in src.lower() or 'APODERADO' in src
    
    # Eliminar párrafos vacíos iniciales y asignar margen exacto de cabecera directamente en el título para que empiece justo debajo de la puntita azul
    to_remove = []
    for p in doc.paragraphs:
        if 'CONTRATO DE CORRETAJE' in p.text:
            p.paragraph_format.space_before = Pt(135) # Margen exacto desde el borde superior hasta debajo del vértice azul
            p.paragraph_format.space_after = Pt(6)
            break
        to_remove.append(p)
    for p in to_remove:
        p._element.getparent().remove(p._element)
    
    # Procesar párrafos optimizando interlineado y espaciado para que todo quepa perfectamente en UNA SOLA HOJA con firmas holgadas
    for i, p in enumerate(list(doc.paragraphs)):
        process_paragraph_keep_exact_font_and_size(p, is_proxy=is_proxy)
        txt = p.text.strip()
        
        # Optimizar espacio entre párrafos para que el contrato completo encaje en 1 hoja sin apretar las firmas
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.page_break_before = False
        
        if 'CONTRATO DE CORRETAJE' in txt:
            p.paragraph_format.space_before = Pt(135)
            p.paragraph_format.space_after = Pt(6)
        elif 'Para constancia de la aceptación' in txt or 'aceptación del contenido total' in txt:
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
        elif ('CORREDOR' in txt and 'CLIENTE' in txt) or (txt and set(txt) <= {'.', '…', ' ', '\t'} and len(txt) > 10):
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(18 if 'CORREDOR' in txt else 2)
            p.paragraph_format.keep_with_next = True
        elif 'CI:' in txt:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        else:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3.5)
        
    # Procesar tablas
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph_keep_exact_font_and_size(p, is_proxy=is_proxy)
                    
    # Aplicar marca de agua y márgenes seguros
    apply_watermark_and_margins(doc, wm_path)
    
    out_path = os.path.join('estate_management/data/templates', dst)
    doc.save(out_path)
    print('Plantilla procesada (márgenes exactos originales + alineación de firmas + marca de agua):', out_path)
